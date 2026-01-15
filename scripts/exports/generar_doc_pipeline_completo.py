#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Genera documento Word con pipeline MOL COMPLETO"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from pathlib import Path
from datetime import datetime

def add_stage_header(doc, numero, titulo):
    """Agrega encabezado de etapa"""
    p = doc.add_paragraph()
    run = p.add_run(f'ETAPA {numero}: {titulo}')
    run.bold = True
    run.font.size = Pt(14)

def main():
    doc = Document()

    # Título
    title = doc.add_heading('Pipeline MOL - Documentación Completa', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f'Versión: 1.0 | Fecha: {datetime.now().strftime("%d/%m/%Y")}')
    doc.add_paragraph()

    # Resumen ejecutivo
    doc.add_heading('Resumen', level=1)
    p = doc.add_paragraph()
    p.add_run('El pipeline procesa ofertas de empleo en 7 etapas: desde el scraping del portal hasta la categorización de skills para dashboards.').bold = True

    doc.add_paragraph()
    doc.add_paragraph('SCRAPING → LIMPIEZA → NLP → POSTPROC → MATCHING → SKILLS → CATEGORÍAS → DASHBOARD')

    # =========================================================================
    # ETAPA 1
    # =========================================================================
    doc.add_page_break()
    add_stage_header(doc, 1, 'SCRAPING')

    doc.add_paragraph('Qué hace: Descarga ofertas de portales de empleo (Bumeran, ZonaJobs)')
    doc.add_paragraph()

    doc.add_paragraph('Ejemplo de datos descargados:', style='Intense Quote')

    tabla = doc.add_table(rows=5, cols=2)
    tabla.style = 'Table Grid'
    datos = [
        ('ID', '1118028201'),
        ('Título', '671SI Operarios ind. ALIMENTICIA c/ Tit. secundario - pres 31/10...'),
        ('Empresa', 'CONA CONSULTORES EN RRHH'),
        ('Ubicación', 'Tortuguitas, Buenos Aires'),
    ]
    tabla.rows[0].cells[0].text = 'Campo'
    tabla.rows[0].cells[1].text = 'Valor'
    for i, (campo, valor) in enumerate(datos, 1):
        tabla.rows[i].cells[0].text = campo
        tabla.rows[i].cells[1].text = valor

    doc.add_paragraph()
    doc.add_paragraph('Problema: El título tiene mucho "ruido" (códigos, fechas, sucursales)')

    # =========================================================================
    # ETAPA 2
    # =========================================================================
    doc.add_paragraph()
    add_stage_header(doc, 2, 'LIMPIEZA DE TÍTULO')

    doc.add_paragraph('Archivo: limpiar_titulos.py + config/nlp_titulo_limpieza.json')
    doc.add_paragraph()

    doc.add_paragraph('ANTES:', style='Intense Quote')
    doc.add_paragraph('"671SI Operarios ind. ALIMENTICIA c/ Tit. secundario - pres 31/10 de 10 a 1130 NUEVA Suc. SAN ISIDRO"')

    doc.add_paragraph('DESPUÉS:', style='Intense Quote')
    p = doc.add_paragraph()
    p.add_run('"Operarios industria ALIMENTICIA"').bold = True

    doc.add_paragraph()
    doc.add_paragraph('Qué se elimina:')
    doc.add_paragraph('• "671SI" → código interno de la consultora', style='List Bullet')
    doc.add_paragraph('• "c/ Tit. secundario" → requisito (va a otro campo)', style='List Bullet')
    doc.add_paragraph('• "pres 31/10 de 10 a 1130" → fecha/hora de entrevista', style='List Bullet')
    doc.add_paragraph('• "NUEVA Suc. SAN ISIDRO" → sucursal de entrevista', style='List Bullet')
    doc.add_paragraph('• "ind." → se expande a "industria"', style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('Otro ejemplo:')
    doc.add_paragraph('ANTES: "Búsqueda Laboral: Modelista – Vicente López (Florida Oeste)"')
    doc.add_paragraph('DESPUÉS: "Modelista"')

    # =========================================================================
    # ETAPA 3
    # =========================================================================
    doc.add_page_break()
    add_stage_header(doc, 3, 'EXTRACCIÓN NLP (Inteligencia Artificial)')

    doc.add_paragraph('Archivo: process_nlp_from_db_v10.py')
    doc.add_paragraph('Modelo: Qwen2.5 (14B parámetros)')
    doc.add_paragraph()

    doc.add_paragraph('Lee la descripción completa y extrae ~50 campos estructurados:')
    doc.add_paragraph()

    tabla = doc.add_table(rows=8, cols=2)
    tabla.style = 'Table Grid'
    tabla.rows[0].cells[0].text = 'Categoría'
    tabla.rows[0].cells[1].text = 'Campos'

    campos_nlp = [
        ('Ubicación', 'provincia, localidad, modalidad (remoto/presencial/híbrido)'),
        ('Requisitos', 'experiencia_min/max, nivel_educativo, idiomas, licencia_conducir'),
        ('Persona', 'requisito_sexo, requisito_edad_min/max'),
        ('Trabajo', 'jornada_laboral, turnos_rotativos, tiene_gente_cargo'),
        ('Salario', 'salario_min/max, moneda, beneficios'),
        ('Skills', 'skills_tecnicas, soft_skills, certificaciones'),
        ('Tareas', 'tareas_explicitas (lista de responsabilidades)'),
    ]
    for i, (cat, campos) in enumerate(campos_nlp, 1):
        tabla.rows[i].cells[0].text = cat
        tabla.rows[i].cells[1].text = campos

    # =========================================================================
    # ETAPA 4
    # =========================================================================
    doc.add_paragraph()
    add_stage_header(doc, 4, 'POSTPROCESSING (Correcciones)')

    doc.add_paragraph('Archivo: nlp_postprocessor.py + configs JSON')
    doc.add_paragraph()
    doc.add_paragraph('Corrige errores del LLM:')
    doc.add_paragraph()

    doc.add_paragraph('1. Ubicación - prioriza dato del scraping:', style='List Bullet')
    doc.add_paragraph('   LLM extrajo: "San Isidro" (lugar de entrevista)')
    doc.add_paragraph('   Scraping dice: "Tortuguitas" (lugar de trabajo real)')
    doc.add_paragraph('   Resultado: "Tortuguitas" ✓')

    doc.add_paragraph()
    doc.add_paragraph('2. Sector - corrige clasificaciones erróneas:', style='List Bullet')
    doc.add_paragraph('   "Modelista" clasificado como "Tecnología" → corregido a "Textil/Confección"')

    doc.add_paragraph()
    doc.add_paragraph('3. Merge skills: combina las extraídas por regex + LLM', style='List Bullet')

    # =========================================================================
    # ETAPA 5
    # =========================================================================
    doc.add_page_break()
    add_stage_header(doc, 5, 'MATCHING ESCO (Clasificación de Ocupación)')

    doc.add_paragraph('Archivo: match_ofertas_v3.py')
    doc.add_paragraph()

    doc.add_paragraph('Entrada: titulo_limpio + tareas + area_funcional + seniority')
    doc.add_paragraph()

    doc.add_paragraph('Proceso de matching:')
    doc.add_paragraph('• Embeddings semánticos del título (modelo BGE-M3) → 50% peso', style='List Bullet')
    doc.add_paragraph('• Búsqueda por skills extraídas → 40% peso', style='List Bullet')
    doc.add_paragraph('• Filtros contextuales (área, seniority) → 10% peso', style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('Ejemplos de matching:', style='Intense Quote')

    tabla = doc.add_table(rows=4, cols=3)
    tabla.style = 'Table Grid'
    tabla.rows[0].cells[0].text = 'Título Limpio'
    tabla.rows[0].cells[1].text = 'ISCO'
    tabla.rows[0].cells[2].text = 'Ocupación ESCO'

    matches = [
        ('Operarios industria ALIMENTICIA', '7511', 'Operario procesamiento alimentos'),
        ('Modelista', '7531', 'Modista'),
        ('Farmacéutico/a', '2262', 'Farmacéutico'),
    ]
    for i, (titulo, isco, esco) in enumerate(matches, 1):
        tabla.rows[i].cells[0].text = titulo
        tabla.rows[i].cells[1].text = isco
        tabla.rows[i].cells[2].text = esco

    # =========================================================================
    # ETAPA 6
    # =========================================================================
    doc.add_paragraph()
    add_stage_header(doc, 6, 'EXTRACCIÓN DE SKILLS IMPLÍCITAS')

    doc.add_paragraph('Archivo: skills_implicit_extractor.py')
    doc.add_paragraph()

    doc.add_paragraph('Lee: titulo_limpio + tareas_explicitas')
    doc.add_paragraph('Busca en: catálogo ESCO con 14,000+ skills usando embeddings')
    doc.add_paragraph()

    doc.add_paragraph('Ejemplo para "Operarios industria ALIMENTICIA":', style='Intense Quote')

    tabla = doc.add_table(rows=6, cols=2)
    tabla.style = 'Table Grid'
    tabla.rows[0].cells[0].text = 'Skill encontrada'
    tabla.rows[0].cells[1].text = 'Score'

    skills = [
        ('manipulación de alimentos', '0.82'),
        ('control de calidad', '0.75'),
        ('normas de higiene', '0.71'),
        ('trabajo en equipo', '0.68'),
        ('BPM (Buenas Prácticas)', '0.65'),
    ]
    for i, (skill, score) in enumerate(skills, 1):
        tabla.rows[i].cells[0].text = skill
        tabla.rows[i].cells[1].text = score

    doc.add_paragraph()
    doc.add_paragraph('Promedio: ~14 skills implícitas por oferta')

    # =========================================================================
    # ETAPA 7
    # =========================================================================
    doc.add_page_break()
    add_stage_header(doc, 7, 'CATEGORIZACIÓN L1/L2 (Para Dashboard)')

    doc.add_paragraph('Archivo: skill_categorizer.py + config/skill_categories.json')
    doc.add_paragraph()

    doc.add_paragraph('Agrupa cada skill en categorías para visualización agregada:')
    doc.add_paragraph()

    doc.add_paragraph('NIVEL L1 - 7 Grupos Principales:', style='Intense Quote')

    tabla = doc.add_table(rows=8, cols=3)
    tabla.style = 'Table Grid'
    tabla.rows[0].cells[0].text = 'Código'
    tabla.rows[0].cells[1].text = 'Nombre'
    tabla.rows[0].cells[2].text = 'Ejemplos'

    categorias = [
        ('S1', 'Comunicación', 'redacción, presentaciones'),
        ('S3', 'Asistencia/Ventas', 'atención al cliente, negociación'),
        ('S4', 'Gestión', 'contabilidad, planificación, presupuestos'),
        ('S5', 'Digital/IT', 'Python, Excel, SAP, SQL'),
        ('S6', 'Técnicas', 'soldadura, electricidad, CNC'),
        ('K', 'Conocimientos', 'normativa legal, inglés técnico'),
        ('T', 'Transversales', 'liderazgo, trabajo en equipo'),
    ]
    for i, (cod, nombre, ejemplos) in enumerate(categorias, 1):
        tabla.rows[i].cells[0].text = cod
        tabla.rows[i].cells[1].text = nombre
        tabla.rows[i].cells[2].text = ejemplos

    doc.add_paragraph()
    doc.add_paragraph('Además: flag "es_digital" (Sí/No) para análisis de brecha digital')

    doc.add_paragraph()
    doc.add_paragraph('Ejemplo de categorización:', style='Intense Quote')
    doc.add_paragraph('"manipulación de alimentos" → L1: S6 (Técnicas), es_digital: No')
    doc.add_paragraph('"Excel avanzado" → L1: S5 (Digital), es_digital: Sí')

    # =========================================================================
    # SALIDA FINAL
    # =========================================================================
    doc.add_paragraph()
    doc.add_heading('Salida Final para Dashboard', level=1)

    doc.add_paragraph('Por cada oferta el sistema entrega:')
    doc.add_paragraph()

    doc.add_paragraph('DATOS NLP:', style='List Bullet')
    doc.add_paragraph('   titulo_limpio, provincia, localidad, sector, modalidad, experiencia, educación, sexo, edad, salario, skills_tecnicas, soft_skills, tareas')

    doc.add_paragraph('MATCHING ESCO:', style='List Bullet')
    doc.add_paragraph('   Código ISCO (ej: 7511) + Ocupación (ej: "Operario procesamiento alimentos")')

    doc.add_paragraph('SKILLS CATEGORIZADAS:', style='List Bullet')
    doc.add_paragraph('   ~14 skills implícitas con categoría L1/L2 + conteo por grupo + % digitales')

    doc.add_paragraph()
    doc.add_paragraph('Esto permite en el Dashboard:')
    doc.add_paragraph('• Distribución de ocupaciones por código ISCO', style='List Bullet')
    doc.add_paragraph('• Mapa de skills más demandadas por sector/región', style='List Bullet')
    doc.add_paragraph('• Análisis de brecha digital', style='List Bullet')
    doc.add_paragraph('• Tendencias temporales del mercado laboral', style='List Bullet')

    # =========================================================================
    # MÉTRICAS
    # =========================================================================
    doc.add_paragraph()
    doc.add_heading('Métricas Actuales', level=1)

    tabla = doc.add_table(rows=5, cols=2)
    tabla.style = 'Table Grid'
    tabla.rows[0].cells[0].text = 'Métrica'
    tabla.rows[0].cells[1].text = 'Valor'

    metricas = [
        ('Ofertas en base de datos', '~10,000'),
        ('Precisión NLP', '~90%'),
        ('Precisión Matching ESCO', '100% (Gold Set 49 ofertas)'),
        ('Skills promedio por oferta', '14'),
    ]
    for i, (m, v) in enumerate(metricas, 1):
        tabla.rows[i].cells[0].text = m
        tabla.rows[i].cells[1].text = v

    # Guardar
    output_path = Path(__file__).parent.parent / "exports" / "Pipeline_MOL_Completo.docx"
    doc.save(str(output_path))
    print(f"Documento guardado en: {output_path}")

if __name__ == "__main__":
    main()
