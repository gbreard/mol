#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Genera documento Word explicando el pipeline MOL"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
from datetime import datetime

def main():
    doc = Document()

    # Título
    title = doc.add_heading('Pipeline MOL - Cómo funciona', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f'Fecha: {datetime.now().strftime("%d/%m/%Y")}')
    doc.add_paragraph()

    # Resumen
    doc.add_heading('Resumen', level=1)
    p = doc.add_paragraph()
    p.add_run('Scrapeamos ofertas de empleo → Limpiamos el texto → Extraemos información con IA → Clasificamos la ocupación según ESCO').bold = True

    # Las 4 etapas
    doc.add_heading('Las 4 Etapas', level=1)

    tabla_etapas = doc.add_table(rows=5, cols=3)
    tabla_etapas.style = 'Table Grid'

    # Headers
    tabla_etapas.rows[0].cells[0].text = 'Etapa'
    tabla_etapas.rows[0].cells[1].text = 'Qué hace'
    tabla_etapas.rows[0].cells[2].text = 'Resultado'

    # Datos
    etapas = [
        ('1. Scraping', 'Descarga ofertas de Bumeran', 'Título + Descripción + Ubicación'),
        ('2. Limpieza', 'Quita ruido del título', 'Título limpio (solo ocupación)'),
        ('3. NLP (IA)', 'Extrae datos estructurados', '50+ campos (experiencia, skills, etc)'),
        ('4. Matching', 'Clasifica ocupación ESCO', 'Código ISCO + nombre ocupación'),
    ]

    for i, (etapa, que_hace, resultado) in enumerate(etapas, 1):
        tabla_etapas.rows[i].cells[0].text = etapa
        tabla_etapas.rows[i].cells[1].text = que_hace
        tabla_etapas.rows[i].cells[2].text = resultado

    doc.add_paragraph()

    # Ejemplo 1: Limpieza
    doc.add_heading('Ejemplo 1: Limpieza de Título', level=1)

    doc.add_paragraph('ANTES (título original del portal):')
    p = doc.add_paragraph()
    p.add_run('"671SI Operarios ind. ALIMENTICIA c/ Tit. secundario - pres 31/10 de 10 a 1130 NUEVA Suc. SAN ISIDRO"').italic = True

    doc.add_paragraph()
    doc.add_paragraph('DESPUÉS (título limpio):')
    p = doc.add_paragraph()
    p.add_run('"Operarios industria ALIMENTICIA"').bold = True

    doc.add_paragraph()
    doc.add_paragraph('¿Qué se eliminó?')
    doc.add_paragraph('• "671SI" → código interno de la consultora', style='List Bullet')
    doc.add_paragraph('• "c/ Tit. secundario" → requisito (va a otro campo)', style='List Bullet')
    doc.add_paragraph('• "pres 31/10 de 10 a 1130" → fecha de entrevista', style='List Bullet')
    doc.add_paragraph('• "NUEVA Suc. SAN ISIDRO" → sucursal de entrevista', style='List Bullet')
    doc.add_paragraph('• "ind." → se expandió a "industria"', style='List Bullet')

    # Ejemplo 2: NLP
    doc.add_heading('Ejemplo 2: Extracción NLP', level=1)

    doc.add_paragraph('De la descripción "Somos Farmacity... buscamos farmacéuticos en Río Cuarto... dispensa de medicamentos, control de stock..."')
    doc.add_paragraph()
    doc.add_paragraph('La IA extrae:')

    tabla_nlp = doc.add_table(rows=8, cols=2)
    tabla_nlp.style = 'Table Grid'

    campos = [
        ('provincia', 'Córdoba'),
        ('localidad', 'Río Cuarto'),
        ('sector_empresa', 'Salud/Farmaceutica'),
        ('nivel_educativo', 'universitario'),
        ('modalidad', 'presencial'),
        ('area_funcional', 'Salud'),
        ('tareas', 'dispensa de medicamentos; control de stock; campañas vacunación'),
    ]

    tabla_nlp.rows[0].cells[0].text = 'Campo'
    tabla_nlp.rows[0].cells[1].text = 'Valor extraído'

    for i, (campo, valor) in enumerate(campos, 1):
        tabla_nlp.rows[i].cells[0].text = campo
        tabla_nlp.rows[i].cells[1].text = valor

    # Ejemplo 3: Matching
    doc.add_heading('Ejemplo 3: Matching ESCO', level=1)

    doc.add_paragraph('El sistema busca la ocupación más parecida en el catálogo europeo ESCO:')
    doc.add_paragraph()

    tabla_match = doc.add_table(rows=4, cols=3)
    tabla_match.style = 'Table Grid'

    tabla_match.rows[0].cells[0].text = 'Título limpio'
    tabla_match.rows[0].cells[1].text = 'Código ISCO'
    tabla_match.rows[0].cells[2].text = 'Ocupación ESCO'

    matches = [
        ('Operarios industria ALIMENTICIA', '7511', 'Operario procesamiento alimentos'),
        ('Modelista', '7531', 'Modista'),
        ('Farmacéutico/a', '2262', 'Farmacéutico'),
    ]

    for i, (titulo, isco, esco) in enumerate(matches, 1):
        tabla_match.rows[i].cells[0].text = titulo
        tabla_match.rows[i].cells[1].text = isco
        tabla_match.rows[i].cells[2].text = esco

    # Campos principales
    doc.add_heading('Campos Principales Extraídos', level=1)

    doc.add_paragraph('Ubicación: provincia, localidad, modalidad (remoto/presencial)', style='List Bullet')
    doc.add_paragraph('Requisitos: experiencia, nivel educativo, idiomas, licencia conducir', style='List Bullet')
    doc.add_paragraph('Persona: sexo requerido, edad mínima/máxima', style='List Bullet')
    doc.add_paragraph('Trabajo: jornada, turnos rotativos, gente a cargo', style='List Bullet')
    doc.add_paragraph('Salario: mínimo, máximo, moneda, beneficios', style='List Bullet')
    doc.add_paragraph('Skills: técnicas, soft skills, certificaciones', style='List Bullet')

    # Métricas
    doc.add_heading('Métricas Actuales', level=1)

    tabla_metricas = doc.add_table(rows=4, cols=2)
    tabla_metricas.style = 'Table Grid'

    metricas = [
        ('Ofertas en base de datos', '~10,000'),
        ('Precisión NLP', '~90%'),
        ('Precisión Matching ESCO', '100% (Gold Set 49 ofertas)'),
    ]

    tabla_metricas.rows[0].cells[0].text = 'Métrica'
    tabla_metricas.rows[0].cells[1].text = 'Valor'

    for i, (metrica, valor) in enumerate(metricas, 1):
        tabla_metricas.rows[i].cells[0].text = metrica
        tabla_metricas.rows[i].cells[1].text = valor

    # Guardar
    output_path = Path(__file__).parent.parent / "exports" / "Pipeline_MOL_Explicacion.docx"
    doc.save(str(output_path))
    print(f"Documento guardado en: {output_path}")

if __name__ == "__main__":
    main()
