"""
Script para convertir PLAN_TECNICO_MOL_v2.0.md a formato Word (.docx)
con formato profesional y traducción al español
"""

import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime
from pathlib import Path

# Diccionario de traducciones (solo títulos y secciones, NO términos técnicos)
TRADUCCIONES = {
    "TECHNICAL PLAN": "PLAN TÉCNICO",
    "TABLE OF CONTENTS": "TABLA DE CONTENIDOS",
    "Architecture and Context": "Arquitectura y Contexto",
    "Current System State": "Estado Actual del Sistema",
    "Identified Gaps": "Brechas Identificadas",
    "Architectural Decisions": "Decisiones Arquitectónicas",
    "Tech Stack": "Stack Tecnológico",
    "Phase 1: ESCO Ontology": "Fase 1: Ontología ESCO",
    "RDF Extraction": "Extracción desde RDF",
    "Knowledge vs Competencies Classification": "Clasificación Conocimientos vs Competencias",
    "Re-matching with Associations": "Re-matching con Asociaciones",
    "Phase 2: Data Pipeline": "Fase 2: Pipeline de Datos",
    "NLP v6.0 - New Fields": "NLP v6.0 - Nuevos Campos",
    "INDEC Territorial Normalization": "Normalización Territorial INDEC",
    "Permanence Calculation": "Cálculo de Permanencia",
    "CSV v2.0 Generation": "Generación CSV v2.0",
    "Phase 3: Shiny Dashboard": "Fase 3: Dashboard Shiny",
    "UI Architecture": "Arquitectura UI",
    "Global Filters": "Filtros Globales",
    "Panel 1: Overview": "Panel 1: Panorama General",
    "Panel 2: Requirements": "Panel 2: Requerimientos",
    "Panel 3: Job Offers": "Panel 3: Ofertas Laborales",
    "Phase 4: Plotly Dashboard v5": "Fase 4: Dashboard Plotly v5",
    "New Tab: Pipeline Monitor": "Nuevo Tab: Pipeline Monitor",
    "Phase 5: Testing and Validation": "Fase 5: Testing y Validación",
    "ESCO Test Suite": "Test Suite ESCO",
    "NLP v6.0 Test Suite": "Test Suite NLP v6.0",
    "Territorial Normalization Test Suite": "Test Suite Normalización Territorial",
    "Quality Validations": "Validaciones de Calidad",
    "Appendices": "Apéndices",
    "Appendix A: Complete SQL Schemas": "Apéndice A: Esquemas SQL Completos",
    "Appendix B: Data Examples": "Apéndice B: Ejemplos de Datos",
    "Appendix C: Configuration": "Apéndice C: Configuración",
    "Appendix D: Troubleshooting": "Apéndice D: Troubleshooting",
    "Appendix E: Useful Commands": "Apéndice E: Comandos Útiles",
}

def crear_estilos_documento(doc):
    """Crear estilos corporativos para el documento"""

    # Estilo para código
    try:
        estilo_codigo = doc.styles.add_style('Codigo', WD_STYLE_TYPE.PARAGRAPH)
    except:
        estilo_codigo = doc.styles['Codigo']

    estilo_codigo.font.name = 'Consolas'
    estilo_codigo.font.size = Pt(9)
    estilo_codigo.paragraph_format.left_indent = Inches(0.5)
    estilo_codigo.paragraph_format.space_before = Pt(6)
    estilo_codigo.paragraph_format.space_after = Pt(6)

    # Configurar estilos de título
    for i in range(1, 5):
        heading = doc.styles[f'Heading {i}']
        heading.font.color.rgb = RGBColor(0, 51, 102)  # Azul oscuro
        heading.font.bold = True

    return doc

def traducir_texto(texto):
    """Traducir solo títulos de secciones, mantener términos técnicos"""
    for eng, esp in TRADUCCIONES.items():
        texto = texto.replace(eng, esp)
    return texto

def agregar_portada(doc):
    """Agregar portada profesional"""
    # Título principal
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run('PLAN TÉCNICO\n')
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)

    run2 = titulo.add_run('Monitor de Ofertas Laborales (MOL) v2.0')
    run2.font.size = Pt(20)
    run2.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph()
    doc.add_paragraph()

    # Metadatos
    metadata = doc.add_paragraph()
    metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fecha_actual = datetime.now().strftime('%B %Y')

    run_meta = metadata.add_run(
        f'Versión: 2.0\n'
        f'Fecha: {fecha_actual}\n'
        f'Autor: Equipo OEDE\n'
        f'Estado: Planificación'
    )
    run_meta.font.size = Pt(12)

    # Salto de página
    doc.add_page_break()

def procesar_linea_markdown(doc, linea):
    """Procesar una línea de Markdown y agregarla al documento Word"""

    linea = linea.rstrip()

    # Omitir líneas vacías múltiples
    if not linea:
        return

    # Traducir la línea
    linea = traducir_texto(linea)

    # Títulos con #
    if linea.startswith('# '):
        texto = linea[2:].strip()
        # Omitir el título principal (ya está en portada)
        if 'PLAN TÉCNICO' not in texto:
            p = doc.add_heading(texto, level=1)
            p.paragraph_format.space_before = Pt(24)
            p.paragraph_format.space_after = Pt(12)
        return

    if linea.startswith('## '):
        texto = linea[3:].strip()
        p = doc.add_heading(texto, level=2)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(10)
        return

    if linea.startswith('### '):
        texto = linea[4:].strip()
        p = doc.add_heading(texto, level=3)
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(8)
        return

    if linea.startswith('#### '):
        texto = linea[5:].strip()
        p = doc.add_heading(texto, level=4)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        return

    # Separadores (---)
    if linea.strip() in ['---', '═══', '───']:
        return

    # Líneas de código con backticks
    if linea.startswith('```'):
        return  # Marcador de inicio/fin de bloque

    # Bullet points
    if linea.startswith('- ') or linea.startswith('* '):
        texto = linea[2:].strip()
        p = doc.add_paragraph(texto, style='List Bullet')
        return

    # Texto normal
    if linea.strip():
        p = doc.add_paragraph(linea.strip())
        p.paragraph_format.space_after = Pt(6)

def procesar_tabla_markdown(doc, lineas_tabla):
    """Procesar tabla en formato Markdown"""
    if not lineas_tabla:
        return

    # Parsear encabezados
    headers = [h.strip() for h in lineas_tabla[0].split('|') if h.strip()]

    # Contar filas (omitir línea separadora)
    filas_datos = [lineas_tabla[i] for i in range(2, len(lineas_tabla))]

    # Crear tabla
    tabla = doc.add_table(rows=1 + len(filas_datos), cols=len(headers))
    tabla.style = 'Light Grid Accent 1'

    # Agregar encabezados
    for i, header in enumerate(headers):
        celda = tabla.rows[0].cells[i]
        celda.text = header
        celda.paragraphs[0].runs[0].font.bold = True

    # Agregar filas de datos
    for idx_fila, fila in enumerate(filas_datos):
        celdas = [c.strip() for c in fila.split('|') if c.strip()]
        for idx_col, contenido in enumerate(celdas):
            if idx_col < len(headers):
                tabla.rows[idx_fila + 1].cells[idx_col].text = contenido

    doc.add_paragraph()  # Espacio después de tabla

def procesar_bloque_codigo(doc, lineas_codigo, lenguaje=''):
    """Procesar bloque de código"""
    if not lineas_codigo:
        return

    # Agregar comentario con lenguaje si existe
    if lenguaje:
        p = doc.add_paragraph(f'[{lenguaje}]', style='Codigo')
        p.runs[0].font.italic = True
        p.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    # Agregar código (sin fondo por limitaciones de python-docx)
    for linea in lineas_codigo:
        p = doc.add_paragraph(linea, style='Codigo')

    doc.add_paragraph()  # Espacio después del código

def convertir_md_a_docx(md_path, docx_path):
    """Función principal de conversión"""

    print(f"[*] Leyendo {md_path}...")
    with open(md_path, 'r', encoding='utf-8') as f:
        lineas = f.readlines()

    print("[*] Creando documento Word...")
    doc = Document()

    # Configurar márgenes
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Crear estilos
    doc = crear_estilos_documento(doc)

    # Agregar portada
    print("[*] Generando portada...")
    agregar_portada(doc)

    # Procesar contenido
    print("[*] Procesando contenido...")

    i = 0
    en_bloque_codigo = False
    lenguaje_codigo = ''
    lineas_codigo = []

    en_tabla = False
    lineas_tabla = []

    while i < len(lineas):
        linea = lineas[i].rstrip()

        # Detectar inicio/fin de bloque de código
        if linea.startswith('```'):
            if not en_bloque_codigo:
                # Inicio de bloque
                en_bloque_codigo = True
                lenguaje_codigo = linea[3:].strip()
                lineas_codigo = []
            else:
                # Fin de bloque
                en_bloque_codigo = False
                procesar_bloque_codigo(doc, lineas_codigo, lenguaje_codigo)
                lineas_codigo = []
                lenguaje_codigo = ''
            i += 1
            continue

        # Si estamos en bloque de código, acumular líneas
        if en_bloque_codigo:
            lineas_codigo.append(linea)
            i += 1
            continue

        # Detectar tablas (líneas con |)
        if '|' in linea and linea.strip():
            if not en_tabla:
                en_tabla = True
                lineas_tabla = [linea]
            else:
                lineas_tabla.append(linea)
            i += 1
            continue
        else:
            # Fin de tabla
            if en_tabla:
                procesar_tabla_markdown(doc, lineas_tabla)
                en_tabla = False
                lineas_tabla = []

        # Procesar línea normal
        procesar_linea_markdown(doc, linea)
        i += 1

    # Agregar encabezado y pie de página
    print("[*] Agregando encabezado y pie de pagina...")

    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = "Plan Técnico MOL v2.0 - OEDE"
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_para.runs[0].font.size = Pt(9)
    header_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = f"Página "
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.runs[0].font.size = Pt(9)

    # Guardar documento
    print(f"[*] Guardando {docx_path}...")
    doc.save(docx_path)

    print(f"[OK] Documento generado exitosamente!")
    print(f"[OK] Archivo: {docx_path}")
    print(f"[OK] Tamaño: {Path(docx_path).stat().st_size / 1024:.1f} KB")

if __name__ == "__main__":
    md_file = Path(__file__).parent.parent / "docs" / "PLAN_TECNICO_MOL_v2.0_EJECUTIVO.md"
    docx_file = Path(__file__).parent.parent / "docs" / "PLAN_TECNICO_MOL_v2.0_EJECUTIVO.docx"

    convertir_md_a_docx(md_file, docx_file)
