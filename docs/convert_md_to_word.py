#!/usr/bin/env python3
"""
Convertidor de Informe Markdown a Word
Genera documento Word con formato profesional
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

def create_styled_document():
    """Crear documento con estilos personalizados"""
    doc = Document()

    # Configurar estilos
    styles = doc.styles

    # Estilo para título principal
    if 'Custom Title' not in styles:
        title_style = styles.add_style('Custom Title', WD_STYLE_TYPE.PARAGRAPH)
        title_font = title_style.font
        title_font.name = 'Calibri'
        title_font.size = Pt(24)
        title_font.bold = True
        title_font.color.rgb = RGBColor(0, 51, 102)  # Azul oscuro
        title_style.paragraph_format.space_after = Pt(12)
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Estilo para subtítulo
    if 'Custom Subtitle' not in styles:
        subtitle_style = styles.add_style('Custom Subtitle', WD_STYLE_TYPE.PARAGRAPH)
        subtitle_font = subtitle_style.font
        subtitle_font.name = 'Calibri'
        subtitle_font.size = Pt(14)
        subtitle_font.color.rgb = RGBColor(68, 68, 68)  # Gris oscuro
        subtitle_style.paragraph_format.space_after = Pt(6)
        subtitle_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Heading 1
    h1_style = styles['Heading 1']
    h1_font = h1_style.font
    h1_font.name = 'Calibri'
    h1_font.size = Pt(18)
    h1_font.bold = True
    h1_font.color.rgb = RGBColor(0, 51, 102)
    h1_style.paragraph_format.space_before = Pt(12)
    h1_style.paragraph_format.space_after = Pt(6)

    # Heading 2
    h2_style = styles['Heading 2']
    h2_font = h2_style.font
    h2_font.name = 'Calibri'
    h2_font.size = Pt(15)
    h2_font.bold = True
    h2_font.color.rgb = RGBColor(31, 78, 120)
    h2_style.paragraph_format.space_before = Pt(10)
    h2_style.paragraph_format.space_after = Pt(4)

    # Heading 3
    h3_style = styles['Heading 3']
    h3_font = h3_style.font
    h3_font.name = 'Calibri'
    h3_font.size = Pt(13)
    h3_font.bold = True
    h3_font.color.rgb = RGBColor(68, 114, 196)
    h3_style.paragraph_format.space_before = Pt(8)
    h3_style.paragraph_format.space_after = Pt(3)

    # Normal
    normal_style = styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Calibri'
    normal_font.size = Pt(11)
    normal_style.paragraph_format.space_after = Pt(6)
    normal_style.paragraph_format.line_spacing = 1.15

    return doc

def add_markdown_content(doc, md_file):
    """Parsear Markdown y agregar al documento Word"""

    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_code_block = False
    in_table = False
    table_lines = []
    skip_next = False

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        if skip_next:
            skip_next = False
            i += 1
            continue

        # Detectar bloque de código
        if line.startswith('```'):
            in_code_block = not in_code_block
            i += 1
            continue

        if in_code_block:
            # Agregar línea de código con formato monospace
            p = doc.add_paragraph(style='Normal')
            p_format = p.paragraph_format
            p_format.left_indent = Inches(0.5)
            if line:  # Solo agregar contenido si la línea no está vacía
                run = p.add_run(line)
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
            i += 1
            continue

        # Detectar tabla
        if '|' in line and line.strip().startswith('|'):
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            i += 1
            continue
        elif in_table and '|' not in line:
            # Fin de tabla, procesarla
            add_table(doc, table_lines)
            in_table = False
            table_lines = []

        # Título principal (# en MD)
        if line.startswith('# ') and not line.startswith('## '):
            text = line[2:].strip()
            if text:
                p = doc.add_paragraph(text, style='Custom Title')

        # Heading 1 (##)
        elif line.startswith('## ') and not line.startswith('### '):
            text = line[3:].strip()
            if text:
                doc.add_paragraph(text, style='Heading 1')

        # Heading 2 (###)
        elif line.startswith('### ') and not line.startswith('#### '):
            text = line[4:].strip()
            if text:
                doc.add_paragraph(text, style='Heading 2')

        # Heading 3 (####)
        elif line.startswith('#### '):
            text = line[5:].strip()
            if text:
                doc.add_paragraph(text, style='Heading 3')

        # Línea horizontal (---)
        elif line.startswith('---') or line.startswith('==='):
            # Agregar separador visual sutil
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)

        # Lista con viñetas (- o *)
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:].strip()
            if text:
                p = doc.add_paragraph(text, style='List Bullet')
                # Procesar negritas
                process_inline_formatting(p)

        # Lista numerada
        elif re.match(r'^\d+\.\s', line.strip()):
            text = re.sub(r'^\d+\.\s', '', line.strip())
            if text:
                p = doc.add_paragraph(text, style='List Number')
                process_inline_formatting(p)

        # Párrafo con formato especial (citas)
        elif line.strip().startswith('>'):
            text = line.strip()[1:].strip()
            if text:
                p = doc.add_paragraph(text, style='Quote')
                p.paragraph_format.left_indent = Inches(0.5)
                run = p.runs[0] if p.runs else None
                if run:
                    run.font.italic = True
                    run.font.color.rgb = RGBColor(100, 100, 100)

        # Línea vacía
        elif not line.strip():
            # Mantener espaciado, no agregar párrafo vacío extra
            pass

        # Párrafo normal
        elif line.strip():
            # Verificar si es metadata (al inicio del documento)
            if i < 10 and line.startswith('**'):
                p = doc.add_paragraph(line, style='Custom Subtitle')
                process_inline_formatting(p)
            else:
                p = doc.add_paragraph(line, style='Normal')
                process_inline_formatting(p)

        i += 1

def process_inline_formatting(paragraph):
    """Procesar negritas, cursivas, etc en un párrafo"""
    # Esta es una versión simplificada
    # En un parser completo, se debería hacer un análisis más sofisticado
    for run in paragraph.runs:
        text = run.text
        # Negritas **texto**
        if '**' in text:
            parts = text.split('**')
            run.text = parts[0]
            for i in range(1, len(parts)):
                if i % 2 == 1:  # Impar = negrita
                    new_run = paragraph.add_run(parts[i])
                    new_run.bold = True
                else:  # Par = normal
                    paragraph.add_run(parts[i])

def add_table(doc, table_lines):
    """Agregar tabla al documento"""
    if len(table_lines) < 2:
        return

    # Filtrar líneas de separador (solo guiones y pipes)
    data_lines = [line for line in table_lines if not re.match(r'^[\|\s\-:]+$', line)]

    if len(data_lines) < 2:
        return

    # Parsear filas
    rows = []
    for line in data_lines:
        cells = [cell.strip() for cell in line.split('|')[1:-1]]  # Excluir pipes al inicio y fin
        if cells:
            rows.append(cells)

    if not rows:
        return

    # Crear tabla
    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Light Grid Accent 1'

    # Llenar celdas
    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j < len(row.cells):
                cell = row.cells[j]
                cell.text = cell_text
                # Header en negrita
                if i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

    # Espaciado después de la tabla
    doc.add_paragraph()

def main():
    print("Convirtiendo Informe Markdown a Word...")
    print()

    md_file = r"D:\OEDE\Webscrapping\docs\INFORME_FUNCIONAL_MOL_NOVIEMBRE_2025.md"
    output_file = r"D:\OEDE\Webscrapping\docs\INFORME_FUNCIONAL_MOL_NOVIEMBRE_2025.docx"

    # Crear documento con estilos
    doc = create_styled_document()

    # Agregar contenido desde Markdown
    print("Parseando Markdown y aplicando estilos...")
    add_markdown_content(doc, md_file)

    # Configurar propiedades del documento
    core_props = doc.core_properties
    core_props.title = "INFORME FUNCIONAL - Monitor de Ofertas Laborales"
    core_props.author = "Observatorio de Empleo y Dinámica Empresarial (OEDE)"
    core_props.subject = "Sistema Integral de Inteligencia del Mercado Laboral Argentino"
    core_props.keywords = "empleo, mercado laboral, NLP, ESCO, intermediación, políticas públicas"

    # Configurar márgenes
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Guardar documento
    print("Guardando documento Word...")
    doc.save(output_file)

    print()
    print("=" * 70)
    print("DOCUMENTO WORD GENERADO EXITOSAMENTE")
    print("=" * 70)
    print()
    print(f"Archivo: {output_file}")
    print()
    print("Características:")
    print("  ✓ Formato profesional con estilos Calibri")
    print("  ✓ Sin saltos de página forzados")
    print("  ✓ Títulos con colores institucionales (azul)")
    print("  ✓ Tablas con formato claro")
    print("  ✓ Párrafos con espaciado adecuado")
    print("  ✓ Márgenes de 1 pulgada")
    print()

if __name__ == '__main__':
    main()
